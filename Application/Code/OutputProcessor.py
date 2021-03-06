from docx import Document
from docx.shared import Cm
import json

from Application.Code import Settings

def buildString(allDangers):
    used = "Name\tSymbols\tHazards\tPrecautions\n"
    for danger in allDangers:
        used += danger[0] + "\t"
        used += str(danger[1]) + "\t"
        for hazard in danger[2]:
            used += hazard[0] + ", "
        if (len(danger[2]) != 0):
            used = used[:-2] + "\t"
        for precaution in danger[3]:
            used += precaution + ", "
        if (len(danger[3]) != 0):
            used = used[:-2] + "\n"
    return used

def doWord(allDangers):

    document = Document()
    summedDangers = mainTable(document, allDangers)
    if(len(summedDangers[0]) != 0):
        document.add_paragraph()
        hazardTable(document, summedDangers[0])
    if(len(summedDangers[1]) != 0):
        document.add_paragraph()
        precautionTable(document, summedDangers[1])

    location = Settings.path

    document.save(location)
    return location

def mainTable(document, allDangers):

    summedHazards = []
    summedPrecautions = []

    table = document.add_table(rows=0, cols=4, style='TableGrid')

    table.columns[0].width = Settings.mainTableColWidth1
    table.columns[1].width = Settings.mainTableColWidth2
    table.columns[2].width = Settings.mainTableColWidth3
    table.columns[3].width = Settings.mainTableColWidth4

    table.add_row()
    hdr_cells = table.rows[0].cells

    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Symbols'
    hdr_cells[2].text = 'Hazards'
    hdr_cells[3].text = 'Precautions'

    size = Settings.size

    for danger in allDangers:
        row_cells = table.add_row().cells
        row_cells[0].text = danger[0]

        paragraph = row_cells[1].paragraphs[0]
        run = paragraph.add_run()
        for symbol in danger[1]:
            run.add_picture("Images\\GHS0" + symbol + ".png", width = size*10000, height = size*10000)

        for hazard in danger[2]:
            summedHazards.append(hazard)
            row_cells[2].text = row_cells[2].text + hazard + ", "
        if (len(danger[2]) != 0):
            row_cells[2].text = row_cells[2].text[:-2]

        for precaution in danger[3]:
            summedPrecautions.append(precaution)
            row_cells[3].text = row_cells[3].text + precaution + ", "
        if (len(danger[3]) != 0):
            row_cells[3].text = row_cells[3].text[:-2]

    return [summedHazards, summedPrecautions]

def hazardTable(document, summedHazards):
    table = document.add_table(rows=0, cols=2, style='TableGrid')
    table.columns[0].width = Settings.discTableColWidth1
    table.columns[1].width = Settings.discTableColWidth2

    with open('HP\\hazards.txt', 'r', encoding='utf-8') as allHazardsFile:
        allHazards = allHazardsFile.read()
        hazardDict = json.loads(allHazards)
        for hazard in hazardDict:
            if(hazard in summedHazards):
                row_cells = table.add_row().cells
                row_cells[0].text = hazard
                row_cells[1].text = hazardDict.get(hazard)

def precautionTable(document, summedPrecautions):
    table = document.add_table(rows=0, cols=2, style='TableGrid')
    table.columns[0].width = Settings.discTableColWidth1
    table.columns[1].width = Settings.discTableColWidth2

    with open('HP\\precautions.txt', 'r', encoding='utf-8') as allPrecautionsFile:
        allPrecautions = allPrecautionsFile.read()
        precautionDict = json.loads(allPrecautions)
        for precaution in precautionDict:
            if(precaution in summedPrecautions):
                row_cells = table.add_row().cells
                row_cells[0].text = precaution
                row_cells[1].text = precautionDict.get(precaution)
